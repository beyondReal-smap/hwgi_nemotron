"""업로드된 파일에서 텍스트 추출.

지원 형식:
- TXT (text/plain) — UTF-8/CP949 자동 인지
- PDF (application/pdf) — pypdf
- DOCX (application/vnd.openxmlformats-officedocument.wordprocessingml.document) — python-docx
- HWP (application/x-hwp) — pyhwp의 hwp5txt CLI (subprocess + 임시파일)
- HWPX (application/vnd.hancom.hwpx) — zipfile + ElementTree (표준 라이브러리)

미지원: RTF/ODT (v2).
"""

from __future__ import annotations

import io
import os
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import docx
from pypdf import PdfReader

# 파일 크기·텍스트 길이 제한
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_OUTPUT_CHARS = 20_000          # AnalyzeRequest와 동일

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx", ".hwp", ".hwpx"}

# HWP 변환 타임아웃 (큰 파일 대비)
HWP_TIMEOUT_SECONDS = 30


class UnsupportedFormatError(Exception):
    """확장자가 화이트리스트에 없음."""


class FileTooLargeError(Exception):
    """파일 크기 초과."""


class TextExtractionError(Exception):
    """파싱 실패 (손상된 파일, 비어있는 결과 등)."""


def extract_from_bytes(content: bytes, filename: str) -> str:
    """파일 내용 바이트 + 파일명 → 추출된 텍스트.

    Raises:
        UnsupportedFormatError: 확장자가 지원되지 않음.
        FileTooLargeError: 크기 초과.
        TextExtractionError: 파싱 실패 또는 빈 결과.
    """
    if len(content) > MAX_FILE_BYTES:
        raise FileTooLargeError(
            f"파일이 {MAX_FILE_BYTES / 1024 / 1024:.0f}MB를 초과합니다 (실제: {len(content) / 1024 / 1024:.1f}MB)"
        )

    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"지원하지 않는 형식: {ext}. 허용: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if ext == ".txt":
        text = _extract_txt(content)
    elif ext == ".pdf":
        text = _extract_pdf(content)
    elif ext == ".docx":
        text = _extract_docx(content)
    elif ext == ".hwp":
        text = _extract_hwp(content)
    elif ext == ".hwpx":
        text = _extract_hwpx(content)
    else:
        # 이미 위에서 걸렀지만 방어
        raise UnsupportedFormatError(f"지원하지 않는 형식: {ext}")

    text = text.strip()
    if not text:
        raise TextExtractionError("파일에서 텍스트를 찾지 못했습니다 (이미지 PDF 등은 OCR이 필요합니다)")

    return text[:MAX_OUTPUT_CHARS]


def _extract_txt(content: bytes) -> str:
    """UTF-8 우선, 실패 시 CP949(EUC-KR) 시도."""
    for encoding in ("utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TextExtractionError("텍스트 파일 인코딩 인식 실패 (UTF-8/CP949 모두 실패)")


def _extract_pdf(content: bytes) -> str:
    """pypdf로 PDF 페이지별 텍스트 추출 + 합치기."""
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as e:
        raise TextExtractionError(f"PDF 파싱 실패: {e}") from e

    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            # 개별 페이지 손상 시 해당 페이지만 예외 격리
            page_text = ""
        if page_text.strip():
            pages.append(page_text)

    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    """python-docx로 단락·표 텍스트 추출."""
    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception as e:
        raise TextExtractionError(f"DOCX 파싱 실패: {e}") from e

    parts: list[str] = []
    # 본문 단락
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 표 셀
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_hwp(content: bytes) -> str:
    """HWP 5.x (OLE 컨테이너) — pyhwp의 hwp5txt CLI 호출.

    OLE 구조 + 압축 레코드 직접 파싱은 복잡하므로 검증된 hwp5txt를 subprocess로 사용.
    임시 파일로 디스크에 쓰고 변환 후 즉시 삭제.
    """
    # 동일 venv의 hwp5txt를 우선 (시스템 PATH 의존 회피)
    hwp5txt = Path(sys.executable).parent / "hwp5txt"
    cmd = str(hwp5txt) if hwp5txt.exists() else "hwp5txt"

    hwp_file = tempfile.NamedTemporaryFile(suffix=".hwp", delete=False)
    hwp_path = hwp_file.name
    try:
        hwp_file.write(content)
        hwp_file.close()  # subprocess가 읽을 수 있도록 확실히 디스크 스트림 닫음
    except Exception as e:
        try:
            os.unlink(hwp_path)
        except OSError:
            pass
        raise TextExtractionError(f"임시 파일 쓰기 실패: {e}") from e

    try:
        result = subprocess.run(
            [cmd, hwp_path],
            capture_output=True,
            timeout=HWP_TIMEOUT_SECONDS,
            check=False,
        )
    except FileNotFoundError as e:
        raise TextExtractionError(
            "hwp5txt 도구를 찾을 수 없습니다 (pyhwp 설치 확인 필요)"
        ) from e
    except subprocess.TimeoutExpired as e:
        raise TextExtractionError(
            f"HWP 변환 시간 초과 ({HWP_TIMEOUT_SECONDS}초)"
        ) from e
    finally:
        try:
            os.unlink(hwp_path)
        except OSError:
            pass

    if result.returncode != 0:
        err = result.stderr.decode("utf-8", errors="replace")[:300]
        raise TextExtractionError(f"HWP 변환 실패: {err.strip() or '알 수 없는 오류'}")

    # hwp5txt는 UTF-8로 stdout 출력
    return result.stdout.decode("utf-8", errors="replace")


# HWPX XML 네임스페이스
_HWPX_NAMESPACES = {
    "hp": "http://www.hancom.co.kr/hwpml/2011/paragraph",
    "hh": "http://www.hancom.co.kr/hwpml/2011/head",
    "hc": "http://www.hancom.co.kr/hwpml/2011/core",
    "hs": "http://www.hancom.co.kr/hwpml/2011/section",
}


def _extract_hwpx(content: bytes) -> str:
    """HWPX (ZIP+XML) — 표준 라이브러리만 사용.

    section*.xml 안의 <hp:t> (text run) 요소들을 모두 추출하여 합침.
    """
    parts: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            # 본문 섹션 파일들 (정렬된 순서로)
            section_names = sorted(
                n for n in zf.namelist()
                if n.startswith("Contents/section") and n.endswith(".xml")
            )

            if not section_names:
                raise TextExtractionError("HWPX 본문 섹션(section*.xml)을 찾지 못했습니다")

            for section_name in section_names:
                try:
                    xml_bytes = zf.read(section_name)
                    root = ET.fromstring(xml_bytes)
                except (KeyError, ET.ParseError) as e:
                    raise TextExtractionError(f"{section_name} 파싱 실패: {e}") from e

                # <hp:t> 또는 t 요소 모두 매칭 (네임스페이스 유무)
                # iter()는 후손 전체를 순회
                for elem in root.iter():
                    tag = elem.tag
                    # 네임스페이스 prefix 제거: {ns}t → t
                    local = tag.rsplit("}", 1)[-1] if "}" in tag else tag
                    if local == "t" and elem.text:
                        parts.append(elem.text)

                # 단락 단위 줄바꿈을 위해 섹션 사이 빈 줄
                parts.append("")
    except zipfile.BadZipFile as e:
        raise TextExtractionError(f"HWPX 파일이 손상되었거나 ZIP 형식이 아닙니다: {e}") from e
    except Exception as e:
        raise TextExtractionError(f"HWPX 텍스트 추출 중 알 수 없는 오류 발생: {e}") from e

    return "\n".join(p for p in parts if p is not None).strip()
