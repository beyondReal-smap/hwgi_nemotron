"""설문 질문 파일 파서 — xlsx/xls/csv/docx.

스펙 §5.5 구현. 파일 → list[ParsedQuestion] 변환.

Excel/CSV 컬럼 동의어:
  - 순서: order, 순번, 번호
  - 유형: type, 유형, 질문유형, 종류
  - 내용: question, text, 질문, 질문내용, 내용
  - 선택지: option1, options, 선택지1, 선택지N (동적)
  - 필수: required, 필수, 필수여부
  - 분기: branch, 분기조건 (현재는 무시, MVP 외)

Word 파싱:
  - 번호 매김(1., Q1, 문1) 정규식으로 질문 시작 감지
  - 다음 질문 전까지의 불릿/체크박스는 선택지
  - 유형 자동 추정: 선택지 없음=주관식 / 1개=척도 / 2개+=단일선택
"""

from __future__ import annotations

import csv
import io
import re
from typing import Literal

from openpyxl import load_workbook
from docx import Document
from pydantic import BaseModel, Field


QuestionType = Literal["single_choice", "multi_choice", "scale", "open_ended", "nps"]


class ParsedQuestion(BaseModel):
    row: int                                # 원본 행 번호(1-based, 사용자 표시용)
    type: QuestionType
    text: str = ""
    options: list[str] = Field(default_factory=list)
    scale_min: int | None = None
    scale_max: int | None = None
    required: bool = True
    errors: list[str] = Field(default_factory=list)  # 행별 검증 오류 메시지


class ParseSummary(BaseModel):
    total: int
    valid: int
    invalid: int


class ParseResult(BaseModel):
    filename: str
    file_format: Literal["xlsx", "csv", "docx"]
    summary: ParseSummary
    questions: list[ParsedQuestion]


# ============================================================
# 동의어 사전
# ============================================================

# 키 = canonical, value = 동의어 list (소문자)
_HEADER_ALIASES: dict[str, list[str]] = {
    "order": ["order", "순번", "번호", "순서", "no", "#"],
    "type": ["type", "유형", "질문유형", "종류", "타입"],
    "text": ["question", "text", "질문", "질문내용", "내용", "본문"],
    "required": ["required", "필수", "필수여부"],
}

_TYPE_ALIASES: dict[str, QuestionType] = {
    "single_choice": "single_choice",
    "single": "single_choice",
    "단일": "single_choice",
    "단일선택": "single_choice",
    "단일 선택": "single_choice",
    "객관식": "single_choice",
    "multi_choice": "multi_choice",
    "multi": "multi_choice",
    "다중": "multi_choice",
    "다중선택": "multi_choice",
    "다중 선택": "multi_choice",
    "scale": "scale",
    "척도": "scale",
    "리커트": "scale",
    "open_ended": "open_ended",
    "open": "open_ended",
    "주관식": "open_ended",
    "자유": "open_ended",
    "nps": "nps",
}


def _canonical_header(s: str) -> str | None:
    """헤더 셀 텍스트 → canonical key. 모르면 None."""
    if not s:
        return None
    norm = s.strip().lower()
    for canon, aliases in _HEADER_ALIASES.items():
        if norm in [a.lower() for a in aliases]:
            return canon
    # '선택지N' 또는 'optionN' 패턴
    m = re.match(r"^(선택지|option)\s*(\d+)$", norm)
    if m:
        return f"option_{int(m.group(2))}"
    return None


def _normalize_type(s: str) -> QuestionType | None:
    """type 셀 → canonical QuestionType. 모르면 None."""
    if not s:
        return None
    norm = s.strip().lower()
    return _TYPE_ALIASES.get(norm)


def _normalize_required(v: str | bool | int | None) -> bool:
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return bool(v)
    if not v:
        return True
    s = str(v).strip().lower()
    if s in ("n", "no", "false", "0", "x", "아니오", "아니요", "불필요"):
        return False
    return True


# ============================================================
# Excel (xlsx)
# ============================================================

def parse_xlsx(file_bytes: bytes) -> list[ParsedQuestion]:
    """첫 시트의 첫 행을 헤더로 인식. 빈 행은 스킵."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    ws = wb.active
    if ws is None:
        return []

    rows_iter = ws.iter_rows(values_only=True)
    try:
        header_row = next(rows_iter)
    except StopIteration:
        return []

    col_map = _build_col_map([str(c) if c is not None else "" for c in header_row])
    if not col_map.get("text"):
        # 헤더 미인식 시 빈 결과 + 행 1에 오류
        return [ParsedQuestion(row=1, type="open_ended", errors=["헤더에서 '질문'/'text' 컬럼을 찾을 수 없습니다"])]

    results: list[ParsedQuestion] = []
    for row_idx, row in enumerate(rows_iter, start=2):
        cells = [str(c).strip() if c is not None else "" for c in row]
        if not any(cells):
            continue  # 빈 행
        results.append(_parse_table_row(row_idx, cells, col_map))
    return results


# ============================================================
# CSV
# ============================================================

def parse_csv(file_bytes: bytes) -> list[ParsedQuestion]:
    text = file_bytes.decode("utf-8-sig", errors="replace")  # BOM 제거
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return []

    header = rows[0]
    col_map = _build_col_map(header)
    if not col_map.get("text"):
        return [ParsedQuestion(row=1, type="open_ended", errors=["헤더에서 '질문'/'text' 컬럼을 찾을 수 없습니다"])]

    results: list[ParsedQuestion] = []
    for row_idx, row in enumerate(rows[1:], start=2):
        cells = [c.strip() for c in row]
        if not any(cells):
            continue
        results.append(_parse_table_row(row_idx, cells, col_map))
    return results


# ============================================================
# 공통 — 표 형식(xlsx/csv)
# ============================================================

def _build_col_map(headers: list[str]) -> dict[str, int]:
    """header 행 → {canonical_key: col_index}. 'option_N'은 'options'로 통합."""
    col_map: dict[str, int] = {}
    options_idx: list[int] = []
    for i, h in enumerate(headers):
        canon = _canonical_header(h)
        if canon is None:
            continue
        if canon.startswith("option_"):
            options_idx.append(i)
        else:
            col_map.setdefault(canon, i)
    if options_idx:
        col_map["_options_indices"] = options_idx  # 특수 키 (list 저장)  # type: ignore
    return col_map


def _parse_table_row(row_idx: int, cells: list[str], col_map: dict) -> ParsedQuestion:
    errors: list[str] = []

    def get(key: str) -> str:
        idx = col_map.get(key)
        if idx is None or idx >= len(cells):
            return ""
        return cells[idx]

    text = get("text")
    raw_type = get("type")
    qtype = _normalize_type(raw_type) or "open_ended"
    if raw_type and not _normalize_type(raw_type):
        errors.append(f"알 수 없는 유형 '{raw_type}' — 주관식으로 추정")

    if not text:
        errors.append("질문 내용이 비어 있습니다")

    options: list[str] = []
    opt_indices = col_map.get("_options_indices") or []
    for i in opt_indices:
        if i < len(cells):
            v = cells[i].strip()
            if v:
                options.append(v)

    # 유형별 검증
    if qtype in ("single_choice", "multi_choice"):
        if len(options) < 2:
            errors.append(f"{qtype}은 선택지 2개 이상 필요 (현재 {len(options)}개)")
    if qtype == "nps":
        scale_min, scale_max = 0, 10
    elif qtype == "scale":
        # scale은 옵션 없이 1-5 기본. min/max 컬럼이 있다면 우선 사용 (MVP 단순화: 항상 1-5)
        scale_min, scale_max = 1, 5
    else:
        scale_min, scale_max = None, None

    return ParsedQuestion(
        row=row_idx,
        type=qtype,
        text=text,
        options=options,
        scale_min=scale_min,
        scale_max=scale_max,
        required=_normalize_required(get("required")),
        errors=errors,
    )


# ============================================================
# Word (docx)
# ============================================================

_QUESTION_START_RE = re.compile(r"^\s*(?:Q\.?\s*\d+|문\.?\s*\d+|\d+[\.\)])\s*(.*)", re.IGNORECASE)


def parse_docx(file_bytes: bytes) -> list[ParsedQuestion]:
    """단락 순회 — 번호 패턴으로 질문 시작 감지, 다음까지 불릿은 선택지."""
    doc = Document(io.BytesIO(file_bytes))

    results: list[ParsedQuestion] = []
    current_text: str | None = None
    current_options: list[str] = []
    current_row: int = 0

    def flush():
        nonlocal current_text, current_options
        if current_text is None:
            return
        # 유형 자동 추정
        opts = [o for o in current_options if o]
        if len(opts) == 0:
            qtype: QuestionType = "open_ended"
        elif len(opts) == 1:
            qtype = "scale"
        else:
            qtype = "single_choice"

        errors = []
        if not current_text.strip():
            errors.append("질문 내용이 비어 있습니다")

        results.append(ParsedQuestion(
            row=current_row,
            type=qtype,
            text=current_text.strip(),
            options=opts,
            scale_min=1 if qtype == "scale" else None,
            scale_max=5 if qtype == "scale" else None,
            errors=errors,
        ))
        current_text = None
        current_options = []

    for idx, p in enumerate(doc.paragraphs, start=1):
        line = (p.text or "").strip()
        if not line:
            continue

        m = _QUESTION_START_RE.match(line)
        if m:
            # 이전 질문 마무리
            flush()
            current_row = idx
            current_text = m.group(1).strip() or ""
            current_options = []
        else:
            # 질문 이어가기 또는 선택지
            style_name = (p.style.name or "").lower() if p.style else ""
            is_list = style_name.startswith("list") or line.startswith("-") or line.startswith("•") or line.startswith("·")
            if current_text is not None:
                if is_list:
                    # 불릿 마커 제거
                    opt = re.sub(r"^[\-•·\s]+", "", line).strip()
                    if opt:
                        current_options.append(opt)
                elif not current_text:
                    # 질문 본문이 다음 줄에 있을 수도
                    current_text = line

    flush()
    return results


# ============================================================
# Dispatcher
# ============================================================

def parse_question_file(
    filename: str, content: bytes,
) -> ParseResult:
    name_lower = filename.lower()
    if name_lower.endswith(".xlsx") or name_lower.endswith(".xls"):
        questions = parse_xlsx(content)
        fmt: Literal["xlsx", "csv", "docx"] = "xlsx"
    elif name_lower.endswith(".csv"):
        questions = parse_csv(content)
        fmt = "csv"
    elif name_lower.endswith(".docx"):
        questions = parse_docx(content)
        fmt = "docx"
    else:
        raise ValueError(
            "지원하지 않는 파일 형식입니다 (.xlsx/.csv/.docx만 가능)"
        )

    valid = sum(1 for q in questions if not q.errors)
    return ParseResult(
        filename=filename,
        file_format=fmt,
        summary=ParseSummary(total=len(questions), valid=valid, invalid=len(questions) - valid),
        questions=questions,
    )


# ============================================================
# 표준 템플릿 생성
# ============================================================

EXCEL_HEADERS = [
    "순서", "질문유형", "질문내용",
    "선택지1", "선택지2", "선택지3", "선택지4", "선택지5",
    "필수여부",
]

EXCEL_SAMPLE_ROWS: list[list] = [
    [1, "single_choice", "사내 식당 점심 만족도?", "매우 불만", "불만", "보통", "만족", "매우 만족", "Y"],
    [2, "scale", "워크라이프 밸런스 만족도 (1-5)", "", "", "", "", "", "Y"],
    [3, "nps", "회사를 동료에게 추천할 의향?", "", "", "", "", "", "Y"],
    [4, "open_ended", "개선이 필요한 점 한 가지", "", "", "", "", "", "N"],
    [5, "multi_choice", "중요하게 여기는 가치(복수)", "연봉", "워라밸", "성장", "동료", "안정성", "Y"],
]

EXCEL_TYPE_NOTE = (
    "* 질문유형: single_choice / multi_choice / scale / open_ended / nps "
    "(한글 가능: 단일선택/다중선택/척도/주관식)\n"
    "* 선택지는 single_choice·multi_choice에만 입력 (2개 이상)\n"
    "* scale은 1-5 고정, nps는 0-10 고정\n"
    "* 필수여부: Y/N (기본 Y)"
)


def build_excel_template() -> bytes:
    """xlsx 표준 템플릿 바이트 반환."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "설문 질문"

    header_fill = PatternFill(start_color="D97757", end_color="D97757", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")

    ws.append(EXCEL_HEADERS)
    for col_idx in range(1, len(EXCEL_HEADERS) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for row in EXCEL_SAMPLE_ROWS:
        ws.append(row)

    # 컬럼 폭
    widths = [8, 14, 38, 14, 14, 14, 14, 14, 10]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w

    # 안내 시트
    ws2 = wb.create_sheet("안내")
    for i, line in enumerate(EXCEL_TYPE_NOTE.split("\n"), start=1):
        ws2.cell(row=i, column=1, value=line)
    ws2.column_dimensions["A"].width = 80

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def build_word_template() -> bytes:
    """docx 표준 템플릿 바이트 반환."""
    doc = Document()
    doc.add_heading("PersonaFit 설문 질문 템플릿", level=1)
    doc.add_paragraph(
        "각 질문은 번호(1., 2. ... 또는 Q1, Q2 ...)로 시작하세요. "
        "선택지는 불릿(-, •) 형태로 다음 줄에 작성합니다. "
        "선택지가 없으면 주관식, 1개면 척도, 2개 이상이면 단일 선택으로 자동 인식됩니다."
    )
    doc.add_paragraph()
    doc.add_paragraph("1. 사내 식당 점심 만족도는 어떠신가요?")
    doc.add_paragraph("- 매우 불만", style="List Bullet")
    doc.add_paragraph("- 불만", style="List Bullet")
    doc.add_paragraph("- 보통", style="List Bullet")
    doc.add_paragraph("- 만족", style="List Bullet")
    doc.add_paragraph("- 매우 만족", style="List Bullet")
    doc.add_paragraph()
    doc.add_paragraph("2. 개선이 필요한 점 한 가지를 짧게 작성해주세요.")
    doc.add_paragraph()
    doc.add_paragraph("3. 중요하게 여기는 가치를 모두 선택해주세요.")
    doc.add_paragraph("- 연봉", style="List Bullet")
    doc.add_paragraph("- 워라밸", style="List Bullet")
    doc.add_paragraph("- 성장 기회", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
